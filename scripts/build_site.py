#!/usr/bin/env python3
"""
Build artifacts for the portfolio site:
- data/publications.json (from data/publications.csv)
- data/summary.json (computed totals and counts)
- assets/Dawid-Tobolski-Scientific-CV.pdf (auto-generated scientific CV)
- assets/Dawid-Tobolski-Scientific-CV.docx (optional; generated as well)

Usage:
  python scripts/build_site.py
"""

from __future__ import annotations

import json
import os
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = PROJECT_ROOT / "data"
ASSETS_DIR = PROJECT_ROOT / "assets"

PUBLICATIONS_CSV = DATA_DIR / "publications.csv"
PUBLICATIONS_JSON = DATA_DIR / "publications.json"
SUMMARY_JSON = DATA_DIR / "summary.json"
PROFILE_JSON = DATA_DIR / "profile.json"
METRICS_JSON = DATA_DIR / "metrics.json"

CV_PDF = ASSETS_DIR / "Dawid-Tobolski-Scientific-CV.pdf"
CV_DOCX = ASSETS_DIR / "Dawid-Tobolski-Scientific-CV.docx"

DATE_FMT = "%d.%m.%Y"


def _rl_escape(s: Any) -> str:
    """Escape text for ReportLab Paragraphs (basic XML entities)."""
    try:
        from xml.sax.saxutils import escape
    except Exception:
        return str(s) if s is not None else ""
    return escape(str(s) if s is not None else "")


def _safe_float(x: Any) -> Optional[float]:
    try:
        if pd.isna(x):
            return None
        return float(x)
    except Exception:
        return None


def _safe_int(x: Any) -> Optional[int]:
    try:
        if pd.isna(x):
            return None
        return int(float(x))
    except Exception:
        return None


def _parse_date_ddmmyyyy(s: Any) -> Optional[datetime]:
    if s is None or (isinstance(s, float) and pd.isna(s)):
        return None
    s = str(s).strip()
    if not s:
        return None
    # Allow both DD.MM.YYYY and YYYY-MM-DD
    for fmt in (DATE_FMT, "%Y-%m-%d"):
        try:
            return datetime.strptime(s, fmt)
        except Exception:
            continue
    return None


def load_profile() -> Dict[str, Any]:
    if not PROFILE_JSON.exists():
        raise FileNotFoundError(f"Missing {PROFILE_JSON}")
    return json.loads(PROFILE_JSON.read_text(encoding="utf-8"))


def load_scholar_metrics() -> Dict[str, Any]:
    if not METRICS_JSON.exists():
        return {
            "source": "Google Scholar",
            "author_id": "",
            "profile_url": "",
            "citations_all": "",
            "citations_since_2016": "",
            "h_index_all": "",
            "h_index_since_2016": "",
            "i10_index_all": "",
            "i10_index_since_2016": "",
            "last_updated": "",
            "note": "metrics.json not found",
        }
    return json.loads(METRICS_JSON.read_text(encoding="utf-8"))


def load_publications_df() -> pd.DataFrame:
    if not PUBLICATIONS_CSV.exists():
        raise FileNotFoundError(f"Missing {PUBLICATIONS_CSV}")

    # Robust CSV loading for Polish diacritics / Excel exports.
    # - Prefer UTF-8 (with or without BOM)
    # - Fallback to common Central European encodings (cp1250, ISO-8859-2)
    # - Auto-detect delimiter (comma vs semicolon) via pandas' python engine
    encodings_to_try = ["utf-8-sig", "utf-8", "cp1250", "iso-8859-2", "latin1"]
    last_err: Optional[Exception] = None

    for enc in encodings_to_try:
        try:
            df = pd.read_csv(
                PUBLICATIONS_CSV,
                dtype=str,
                keep_default_na=False,
                encoding=enc,
                sep=None,
                engine="python",
            ).fillna("")
            break
        except Exception as e:
            last_err = e
            df = None  # type: ignore

    if last_err and df is None:  # type: ignore
        raise last_err

    # Coerce numeric columns
    df["year_int"] = df["year"].apply(_safe_int)
    df["mnicsw_points_num"] = df["mnicsw_points"].apply(_safe_float)
    df["impact_factor_num"] = df["impact_factor"].apply(_safe_float)
    # Dates (for conferences)
    df["start_date_dt"] = df["start_date"].apply(_parse_date_ddmmyyyy)
    df["end_date_dt"] = df["end_date"].apply(_parse_date_ddmmyyyy)
    return df


def compute_summary(df: pd.DataFrame, metrics: Dict[str, Any]) -> Dict[str, Any]:
    def count_where(mask) -> int:
        return int(mask.sum())

    total_points = float(pd.Series(df["mnicsw_points_num"]).dropna().sum())
    total_if = float(pd.Series(df["impact_factor_num"]).dropna().sum())

    is_list_a = df["category"].str.upper().eq("A")
    is_chapter = df["record_type"].eq("book_chapter") | df["category"].str.lower().str.contains("book chapter")
    is_list_b = df["category"].str.upper().eq("B") | is_chapter
    is_conf = df["record_type"].eq("conference_contribution") | df["category"].str.lower().eq("conference")

    # Conference breakdown (best-effort)
    conf_oral = is_conf & df["subtype"].str.lower().str.contains("oral")
    conf_poster = is_conf & df["subtype"].str.lower().str.contains("poster")

    # Yearly counts for a small sparkline / overview
    year_counts = (
        df.dropna(subset=["year_int"])
          .groupby("year_int")
          .size()
          .sort_index(ascending=False)
          .to_dict()
    )

    summary = {
        "computed_from": str(PUBLICATIONS_CSV.as_posix()),
        "generated_on": datetime.now().strftime(DATE_FMT),
        "totals": {
            "mnicsw_points": int(total_points) if float(total_points).is_integer() else total_points,
            "sum_impact_factor": round(total_if, 3),
            "records_total": int(len(df)),
            "publications_list_a": count_where(is_list_a),
            "publications_list_b": count_where(is_list_b),
            "book_chapters": count_where(is_chapter),
            "conference_contributions_total": count_where(is_conf),
            "conference_oral_presentations": count_where(conf_oral),
            "conference_posters": count_where(conf_poster),
            "conference_type_unspecified": count_where(is_conf & ~(conf_oral | conf_poster)),
        },
        "scholar_metrics": {
            "citations_all": metrics.get("citations_all", ""),
            "h_index_all": metrics.get("h_index_all", ""),
            "i10_index_all": metrics.get("i10_index_all", ""),
            "last_updated": metrics.get("last_updated", ""),
            "profile_url": metrics.get("profile_url", ""),
        },
        "year_counts": {str(k): int(v) for k, v in year_counts.items()},
    }
    return summary


def export_publications_json(df: pd.DataFrame) -> None:
    # Sort:
    # - Publications: year desc
    # - Conferences with dates: start_date desc
    def sort_key(row) -> Tuple[int, float]:
        y = row.get("year_int")
        y = y if isinstance(y, int) else -1
        # Use start_date for conference if available
        dt = row.get("start_date_dt")
        ts = dt.timestamp() if isinstance(dt, datetime) and not pd.isna(dt) else 0.0
        return (y, ts)

    df_sorted = df.copy()
    df_sorted["_sort"] = df_sorted.apply(sort_key, axis=1)
    df_sorted = df_sorted.sort_values(by="_sort", ascending=False).drop(columns=["_sort"])

    # Keep only user-facing fields (stringify)
    keep_cols = [
        "record_type",
        "year",
        "category",
        "subtype",
        "citation",
        "doi",
        "mnicsw_points",
        "impact_factor",
        "start_date",
        "end_date",
        "city",
        "country",
        "award",
    ]
    out = df_sorted[keep_cols].to_dict(orient="records")
    PUBLICATIONS_JSON.write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8")


def _format_date_range(start: str, end: str) -> str:
    start_dt = _parse_date_ddmmyyyy(start)
    end_dt = _parse_date_ddmmyyyy(end)
    if start_dt and end_dt:
        if start_dt.date() == end_dt.date():
            return start_dt.strftime(DATE_FMT)
        return f"{start_dt.strftime(DATE_FMT)}–{end_dt.strftime(DATE_FMT)}"
    if start_dt:
        return start_dt.strftime(DATE_FMT)
    return ""


def generate_cv_pdf(profile: Dict[str, Any], summary: Dict[str, Any], df: pd.DataFrame) -> None:
    """
    Creates a clean, ATS-friendly PDF CV from publications.csv.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    ASSETS_DIR.mkdir(parents=True, exist_ok=True)

    # Use a Unicode-capable font to correctly render Polish diacritics (ą, ć, ę, ł, ń, ó, ś, ź, ż)
    # and common scientific symbols (α, β, δ, ®, etc.).
    fonts_dir = ASSETS_DIR / "fonts"
    font_regular = fonts_dir / "DejaVuSans.ttf"
    font_bold = fonts_dir / "DejaVuSans-Bold.ttf"

    base_font = "Helvetica"
    bold_font = "Helvetica-Bold"
    try:
        if font_regular.exists():
            pdfmetrics.registerFont(TTFont("DejaVuSans", str(font_regular)))
            base_font = "DejaVuSans"
            if font_bold.exists():
                pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", str(font_bold)))
                bold_font = "DejaVuSans-Bold"
            else:
                bold_font = base_font
    except Exception:
        # Fall back to base-14 fonts (may not support all Unicode glyphs)
        base_font = "Helvetica"
        bold_font = "Helvetica-Bold"

    doc = SimpleDocTemplate(
        str(CV_PDF),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title="Scientific CV",
        author=profile.get("name", ""),
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="H1", parent=styles["Heading1"], fontName=bold_font, fontSize=16, spaceAfter=6))
    styles.add(ParagraphStyle(name="H2", parent=styles["Heading2"], fontName=bold_font, fontSize=12, spaceAfter=4))
    styles.add(ParagraphStyle(name="Body", parent=styles["BodyText"], fontName=base_font, fontSize=9.5, leading=12))
    styles.add(ParagraphStyle(name="Small", parent=styles["BodyText"], fontName=base_font, fontSize=8.8, leading=11))

    story: List[Any] = []

    name = profile.get("name", "Name")
    role = profile.get("role", "")
    affiliation = profile.get("affiliation", "")
    location = profile.get("location", "")
    email = profile.get("email", "")
    links = profile.get("links", {})

    story.append(Paragraph(_rl_escape(name), styles["H1"]))
    header_lines = []
    if role:
        header_lines.append(role)
    if affiliation:
        header_lines.append(affiliation)
    if location:
        header_lines.append(location)
    if email:
        header_lines.append(email)
    story.append(Paragraph(_rl_escape(" · ".join(header_lines)), styles["Body"]))
    if links:
        # Render as plain text; URLs in PDFs can be clickable depending on viewer
        link_line = " · ".join([f"{k}: {v}" for k, v in links.items()])
        story.append(Paragraph(_rl_escape(link_line), styles["Small"]))

    story.append(Spacer(1, 10))

    # Summary table
    totals = summary.get("totals", {})
    scholar = summary.get("scholar_metrics", {})

    summary_rows = [
        ["MNiSW points (computed)", str(totals.get("mnicsw_points", ""))],
        ["Sum of journal impact factors (computed)", str(totals.get("sum_impact_factor", ""))],
        ["Publications (List A)", str(totals.get("publications_list_a", ""))],
        ["Other publications (List B)", str(totals.get("publications_list_b", ""))],
        ["Book chapters", str(totals.get("book_chapters", ""))],
        ["Conference contributions (total)", str(totals.get("conference_contributions_total", ""))],
        ["Google Scholar citations", str(scholar.get("citations_all", ""))],
        ["Google Scholar h-index", str(scholar.get("h_index_all", ""))],
        ["Google Scholar last updated", str(scholar.get("last_updated", ""))],
    ]
    t = Table(summary_rows, colWidths=[65*mm, 110*mm])
    t.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, -1), base_font, 9),
        ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]))
    story.append(Paragraph("Summary", styles["H2"]))
    story.append(t)
    story.append(Spacer(1, 10))

    # Helpers to add sections
    def add_section(title: str, mask: pd.Series) -> None:
        story.append(Paragraph(_rl_escape(title), styles["H2"]))
        sub = df.loc[mask].copy()
        # Sort by year desc
        sub["year_int"] = sub["year"].apply(_safe_int)
        sub = sub.sort_values(by=["year_int"], ascending=False)
        if sub.empty:
            story.append(Paragraph("No records.", styles["Body"]))
            story.append(Spacer(1, 6))
            return

        # Numbered list as paragraphs (ReportLab doesn't do automatic numbering well without ListFlowable)
        for i, row in enumerate(sub.itertuples(index=False), start=1):
            citation = getattr(row, "citation", "")
            doi = getattr(row, "doi", "")
            pts = getattr(row, "mnicsw_points", "")
            if_ = getattr(row, "impact_factor", "")
            extras = []
            if pts:
                extras.append(f"MNiSW: {pts}")
            if if_:
                extras.append(f"IF: {if_}")
            if doi:
                extras.append(f"DOI: {doi}")
            extra_line = f" ({'; '.join(extras)})" if extras else ""
            story.append(Paragraph(_rl_escape(f"{i}. {citation}{extra_line}"), styles["Body"]))
            story.append(Spacer(1, 2))
        story.append(Spacer(1, 8))

    add_section("Peer-reviewed publications (List A)", df["category"].str.upper().eq("A"))
    add_section("Other publications (List B)", df["category"].str.upper().eq("B"))
    add_section("Book chapters", df["record_type"].eq("book_chapter") | df["category"].str.lower().str.contains("book chapter"))

    # Conferences
    story.append(Paragraph(_rl_escape("Conference contributions"), styles["H2"]))
    conf = df[df["record_type"].eq("conference_contribution") | df["category"].str.lower().eq("conference")].copy()
    if conf.empty:
        story.append(Paragraph("No records.", styles["Body"]))
    else:
        conf["year_int"] = conf["year"].apply(_safe_int)
        # Sort by start date if present else year desc
        conf["start_dt"] = conf["start_date"].apply(_parse_date_ddmmyyyy)
        conf["_sort"] = conf.apply(
            lambda r: (r["start_dt"].timestamp() if isinstance(r["start_dt"], datetime) and not pd.isna(r["start_dt"]) else 0.0, r["year_int"] or 0),
            axis=1
        )
        conf = conf.sort_values(by="_sort", ascending=False).drop(columns=["_sort"])
        for i, row in enumerate(conf.itertuples(index=False), start=1):
            citation = getattr(row, "citation", "")
            subtype = getattr(row, "subtype", "")
            city = getattr(row, "city", "")
            country = getattr(row, "country", "")
            award = getattr(row, "award", "")
            dr = _format_date_range(getattr(row, "start_date", ""), getattr(row, "end_date", ""))
            loc = ", ".join([x for x in [city, country] if x])
            tail_parts = [p for p in [subtype, dr, loc, award] if p]
            tail = " — " + " · ".join(tail_parts) if tail_parts else ""
            story.append(Paragraph(_rl_escape(f"{i}. {citation}{tail}"), styles["Body"]))
            story.append(Spacer(1, 2))

    doc.build(story)


def generate_cv_docx(profile: Dict[str, Any], summary: Dict[str, Any], df: pd.DataFrame) -> None:
    """
    Optional DOCX version of the scientific CV.
    """
    try:
        from docx import Document
    except Exception:
        return

    ASSETS_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(profile.get("name", "Name"), level=0)

    header_line = " · ".join([x for x in [
        profile.get("role", ""),
        profile.get("affiliation", ""),
        profile.get("location", ""),
        profile.get("email", ""),
    ] if x])
    if header_line:
        doc.add_paragraph(header_line)

    links = profile.get("links", {})
    if links:
        doc.add_paragraph(" · ".join([f"{k}: {v}" for k, v in links.items()]))

    doc.add_heading("Summary", level=1)
    totals = summary.get("totals", {})
    scholar = summary.get("scholar_metrics", {})
    items = [
        ("MNiSW points (computed)", totals.get("mnicsw_points", "")),
        ("Sum of journal impact factors (computed)", totals.get("sum_impact_factor", "")),
        ("Publications (List A)", totals.get("publications_list_a", "")),
        ("Other publications (List B)", totals.get("publications_list_b", "")),
        ("Book chapters", totals.get("book_chapters", "")),
        ("Conference contributions (total)", totals.get("conference_contributions_total", "")),
        ("Google Scholar citations", scholar.get("citations_all", "")),
        ("Google Scholar h-index", scholar.get("h_index_all", "")),
        ("Google Scholar last updated", scholar.get("last_updated", "")),
    ]
    for k, v in items:
        doc.add_paragraph(f"{k}: {v}")

    def add_section(title: str, mask: pd.Series) -> None:
        doc.add_heading(title, level=1)
        sub = df.loc[mask].copy()
        sub["year_int"] = sub["year"].apply(_safe_int)
        sub = sub.sort_values(by=["year_int"], ascending=False)
        if sub.empty:
            doc.add_paragraph("No records.")
            return
        for i, row in enumerate(sub.itertuples(index=False), start=1):
            citation = getattr(row, "citation", "")
            doi = getattr(row, "doi", "")
            pts = getattr(row, "mnicsw_points", "")
            if_ = getattr(row, "impact_factor", "")
            extras = []
            if pts:
                extras.append(f"MNiSW: {pts}")
            if if_:
                extras.append(f"IF: {if_}")
            if doi:
                extras.append(f"DOI: {doi}")
            extra_line = f" ({'; '.join(extras)})" if extras else ""
            doc.add_paragraph(f"{i}. {citation}{extra_line}")

    add_section("Peer-reviewed publications (List A)", df["category"].str.upper().eq("A"))
    add_section("Other publications (List B)", df["category"].str.upper().eq("B"))
    add_section("Book chapters", df["record_type"].eq("book_chapter") | df["category"].str.lower().str.contains("book chapter"))

    # Conferences
    doc.add_heading("Conference contributions", level=1)
    conf = df[df["record_type"].eq("conference_contribution") | df["category"].str.lower().eq("conference")].copy()
    conf["year_int"] = conf["year"].apply(_safe_int)
    conf["start_dt"] = conf["start_date"].apply(_parse_date_ddmmyyyy)
    conf["_sort"] = conf.apply(
        lambda r: (r["start_dt"].timestamp() if isinstance(r["start_dt"], datetime) and not pd.isna(r["start_dt"]) else 0.0, r["year_int"] or 0),
        axis=1
    )
    conf = conf.sort_values(by="_sort", ascending=False).drop(columns=["_sort"])
    if conf.empty:
        doc.add_paragraph("No records.")
    else:
        for i, row in enumerate(conf.itertuples(index=False), start=1):
            citation = getattr(row, "citation", "")
            subtype = getattr(row, "subtype", "")
            city = getattr(row, "city", "")
            country = getattr(row, "country", "")
            award = getattr(row, "award", "")
            dr = _format_date_range(getattr(row, "start_date", ""), getattr(row, "end_date", ""))
            loc = ", ".join([x for x in [city, country] if x])
            tail_parts = [p for p in [subtype, dr, loc, award] if p]
            tail = " — " + " · ".join(tail_parts) if tail_parts else ""
            doc.add_paragraph(f"{i}. {citation}{tail}")

    doc.save(str(CV_DOCX))


def main() -> None:
    profile = load_profile()
    metrics = load_scholar_metrics()
    df = load_publications_df()

    # Export publications.json
    export_publications_json(df)

    # Summary.json
    summary = compute_summary(df, metrics)
    SUMMARY_JSON.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")

    # CV
    generate_cv_pdf(profile, summary, df)
    generate_cv_docx(profile, summary, df)

    print("Build completed.")
    print(f"- {PUBLICATIONS_JSON}")
    print(f"- {SUMMARY_JSON}")
    print(f"- {CV_PDF}")
    print(f"- {CV_DOCX}")


if __name__ == "__main__":
    main()
